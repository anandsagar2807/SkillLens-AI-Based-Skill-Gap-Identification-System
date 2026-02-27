"""
File Parser Utility
Extracts text from PDF and DOCX files
"""

import io
from typing import Optional

# PDF parsing
import PyPDF2

# DOCX parsing
from docx import Document


class FileParser:
    """Handles text extraction from various file formats"""
    
    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> Optional[str]:
        """
        Extract text from uploaded file
        
        Args:
            file_content: Raw file bytes
            filename: Original filename with extension
            
        Returns:
            Extracted text string or None if extraction fails
        """
        # Determine file type from extension
        extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        try:
            if extension == 'pdf':
                return FileParser._extract_from_pdf(file_content)
            elif extension in ['docx', 'doc']:
                return FileParser._extract_from_docx(file_content)
            else:
                # Try to decode as plain text
                return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Error extracting text from {filename}: {str(e)}")
            return None
    
    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_content: Raw PDF file bytes
            
        Returns:
            Extracted text string
        """
        text = ""
        try:
            # Create PDF reader object
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            # Extract text from each page
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    
        except Exception as e:
            raise Exception(f"PDF extraction error: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def _extract_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_content: Raw DOCX file bytes
            
        Returns:
            Extracted text string
        """
        text = ""
        try:
            # Create Document object
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"DOCX extraction error: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def get_supported_extensions() -> list:
        """Return list of supported file extensions"""
        return ['pdf', 'docx', 'doc', 'txt']


# Convenience functions
def parse_file(file_content: bytes, filename: str) -> Optional[str]:
    """
    Convenience function to parse a file
    
    Args:
        file_content: Raw file bytes
        filename: Original filename
        
    Returns:
        Extracted text or None
    """
    return FileParser.extract_text(file_content, filename)